import streamlit as st
import io
import json
from datetime import datetime

# Try to import docx libraries
try:
    import docx
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

# Try to import PDF libraries
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

# Configure page
st.set_page_config(
    page_title="CoverLet AI",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Check for missing dependencies at startup
if not DOCX_SUPPORT:
    st.error("❌ **Missing Dependency**: python-docx is not installed. Please add it to your dependencies.")
    st.info("💡 **Solution**: Add `python-docx>=0.8.11` to your config.toml dependencies or requirements.txt")
    st.stop()

# Custom CSS with clean dark theme
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
    
    .stApp {
        background: #0f1419;
        color: #e6edf3;
        font-family: 'Inter', sans-serif;
    }
    
    .main .block-container {
        padding: 1rem 2rem;
        max-width: 1200px;
        background: #161b22;
        border-radius: 8px;
        border: 1px solid #30363d;
        margin: 1rem auto;
    }
    
    .main-header {
        font-size: 2.2rem;
        font-weight: 700;
        color: #f0f6fc;
        text-align: center;
        margin-bottom: 1.5rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #238636;
    }
    
    .section-header {
        font-size: 1.1rem;
        font-weight: 600;
        color: #7c3aed;
        margin: 1.5rem 0 0.8rem 0;
        padding: 0.5rem 0;
        border-bottom: 1px solid #30363d;
    }
    
    .stTextArea > div > div > textarea,
    .stTextInput > div > div > input {
        background-color: #21262d !important;
        border: 1px solid #30363d !important;
        color: #ffffff !important;
        border-radius: 6px !important;
        font-family: 'Inter', sans-serif !important;
    }
    
    .stTextArea > div > div > textarea:focus,
    .stTextInput > div > div > input:focus {
        border-color: #7c3aed !important;
        box-shadow: 0 0 0 2px rgba(124, 58, 237, 0.3) !important;
    }
    
    .stFileUploader > div {
        background: #2d1b69 !important;
        border: 1px dashed #7c3aed !important;
        border-radius: 6px !important;
        padding: 1rem !important;
    }
    
    .stFileUploader > div:hover {
        border-color: #8b5cf6 !important;
        background: #3730a3 !important;
    }
    
    .stFileUploader div[data-testid="stFileUploaderDropzone"] {
        background: #2d1b69 !important;
    }
    
    .stFileUploader div[data-testid="stFileUploaderDropzone"] * {
        color: #c4b5fd !important;
        background: transparent !important;
    }
    
    .stFileUploader div[data-testid="stFileUploaderDropzoneInput"] {
        background: #2d1b69 !important;
    }
    
    .stFileUploader small {
        color: #a78bfa !important;
    }
    
    .stFileUploader p {
        color: #c4b5fd !important;
    }
    
    .stFileUploader span {
        color: #c4b5fd !important;
    }
    
    .stFileUploader button {
        background: #7c3aed !important;
        color: #ffffff !important;
        border: none !important;
        border-radius: 4px !important;
        padding: 0.3rem 0.8rem !important;
        font-size: 0.8rem !important;
    }
    
    .stFileUploader button:hover {
        background: #8b5cf6 !important;
    }
    
    /* Override the nuclear white rule for file uploaders specifically */
    .stFileUploader .stApp * {
        color: #c4b5fd !important;
    }
    
    .stRadio > div {
        background: #21262d !important;
        border: 1px solid #30363d !important;
        border-radius: 6px !important;
        padding: 0.8rem !important;
    }
    
    .stRadio label {
        color: #ffffff !important;
    }
    
    .stRadio div[role="radiogroup"] label {
        color: #ffffff !important;
    }
    
    .stRadio div[role="radiogroup"] label span {
        color: #ffffff !important;
    }
    
    .stRadio div[data-testid="stRadio"] > label > div {
        color: #ffffff !important;
    }
    
    .stRadio div[data-testid="stRadio"] > label > div > span {
        color: #ffffff !important;
    }
    
    .stRadio label span {
        color: #ffffff !important;
    }
    
    .stRadio span {
        color: #ffffff !important;
    }
    
    .stRadio p {
        color: #ffffff !important;
    }
    
    .stRadio div {
        color: #ffffff !important;
    }
    
    .stCheckbox > label {
        background: #21262d;
        border: 1px solid #30363d;
        border-radius: 6px;
        padding: 0.5rem;
        color: #ffffff !important;
    }
    
    .stButton > button {
        background: #7c3aed !important;
        color: white !important;
        border: none !important;
        border-radius: 6px !important;
        padding: 0.5rem 1.5rem !important;
        font-weight: 600 !important;
        transition: background 0.2s !important;
    }
    
    .stButton > button:hover {
        background: #8b5cf6 !important;
    }
    
    .stDownloadButton > button {
        background: #238636 !important;
        color: white !important;
        border: none !important;
        border-radius: 6px !important;
        padding: 0.5rem 1.5rem !important;
        font-weight: 600 !important;
    }
    
    .stDownloadButton > button:hover {
        background: #2ea043 !important;
    }
    
    .stSuccess {
        background: #0f2a1a !important;
        border: 1px solid #238636 !important;
        color: #2ea043 !important;
        border-radius: 6px !important;
    }
    
    .stError {
        background: #2a0f0f !important;
        border: 1px solid #da3633 !important;
        color: #f85149 !important;
        border-radius: 6px !important;
    }
    
    .stInfo {
        background: #0f1a2a !important;
        border: 1px solid #1f6feb !important;
        color: #58a6ff !important;
        border-radius: 6px !important;
    }
    
    .streamlit-expanderHeader {
        background: #21262d !important;
        border: 1px solid #30363d !important;
        border-radius: 6px !important;
        color: #ffffff !important;
    }
    
    .stTextArea label,
    .stTextInput label,
    .stFileUploader label {
        color: #ffffff !important;
        font-weight: 500 !important;
    }
    
    .stRadio div[role="radiogroup"] label {
        color: #ffffff !important;
    }
    
    /* Force ALL text to be white - nuclear option */
    .stApp * {
        color: #ffffff !important;
    }
    
    /* Override specific elements that should stay colored */
    .section-header {
        color: #7c3aed !important;
    }
    
    .main-header {
        color: #f0f6fc !important;
    }
    
    .stSuccess * {
        color: #2ea043 !important;
    }
    
    .stError * {
        color: #f85149 !important;
    }
    
    .stInfo * {
        color: #58a6ff !important;
    }
    
    .stButton button {
        color: white !important;
    }
    
    .stDownloadButton button {
        color: white !important;
    }
    
    /* Placeholder text styling */
    .stTextArea > div > div > textarea::placeholder,
    .stTextInput > div > div > input::placeholder {
        color: #8b949e !important;
    }
    
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
    
    .element-container {
        margin-bottom: 0.5rem !important;
    }
    
    .row-widget {
        margin-bottom: 0.5rem !important;
    }
    
    .stTextArea > div, 
    .stTextInput > div {
        margin-bottom: 0.3rem !important;
    }
</style>
""", unsafe_allow_html=True)

# Title
st.markdown('<h1 class="main-header">⚡ CoverLet AI</h1>', unsafe_allow_html=True)

# Show dependency status
if not PDF_SUPPORT:
    st.warning("📄 **PDF Support**: PyPDF2 not available. Upload .docx files only.")

# Initialize session state
if 'default_prompt' not in st.session_state:
    st.session_state.default_prompt = ""
if 'generated_document' not in st.session_state:
    st.session_state.generated_document = None

def read_docx_file(uploaded_file):
    """Extract text from uploaded Word document"""
    if not DOCX_SUPPORT:
        st.error("Word document support not available")
        return None
        
    try:
        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        st.error(f"Error reading Word document: {str(e)}")
        return None

def read_pdf_file(uploaded_file):
    """Extract text from uploaded PDF document"""
    if not PDF_SUPPORT:
        st.error("PDF support not available. Please install PyPDF2: pip install PyPDF2")
        return None
    
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
        text = []
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
        return '\n'.join(text)
    except Exception as e:
        st.error(f"Error reading PDF document: {str(e)}")
        return None

def generate_default_prompt(language_option, custom_prompt, keep_one_page):
    """Generate the default AI prompt based on user selections"""
    
    # Base prompt structure
    base_prompt = "Please edit the below cover letter aligning it with the job description posted right below this paragraph."
    
    # Language/tone options
    if language_option == "Professional Language":
        tone_instruction = "Use professional, formal language with a confident and polished tone."
    elif language_option == "Match Template Language":
        tone_instruction = "Use a tone and language matching the original template."
    else:  # Custom
        tone_instruction = custom_prompt if custom_prompt else "Use appropriate professional language."
    
    # Page length instruction
    length_instruction = " Keep it one page." if keep_one_page else ""
    
    # Additional instructions
    additional_instructions = " Make sure to use accurate statements matching the information from the resume while highlighting the strong alignment to the role. The generated cover letter should capture the interest of the recruiter in my experience and capabilities. Do not make up stuff or hallucinate innacurate inforamtion. Keep it very accurate but powerful. Do not be overenthusiastic, but confident and somewhat persuasive. Please delete this and the job description while applying cover letter modifications."
    
    # Construct prompt
    prompt = f"{base_prompt} {tone_instruction}{length_instruction}{additional_instructions}"
    
    return prompt

def create_word_document(prompt_text, job_description, template_text, resume_text):
    """Create a Word document with prompt, job description, and template"""
    if not DOCX_SUPPORT:
        st.error("Cannot create Word document - python-docx not available")
        return None
        
    doc = docx.Document()
    
    # Add prompt section
    prompt_heading = doc.add_paragraph()
    prompt_heading.add_run("AI PROMPT INSTRUCTIONS:").bold = True
    prompt_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    prompt_para = doc.add_paragraph(prompt_text)
    prompt_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Add spacing
    doc.add_paragraph()
    
    # Add resume information section
    resume_heading = doc.add_paragraph()
    resume_heading.add_run("RESUME INFORMATION:").bold = True
    resume_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    resume_para = doc.add_paragraph(resume_text)
    resume_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Add spacing
    doc.add_paragraph()
    
    # Add job description section
    job_heading = doc.add_paragraph()
    job_heading.add_run("JOB DESCRIPTION:").bold = True
    job_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    job_para = doc.add_paragraph(job_description)
    job_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    # Add spacing and separator
    doc.add_paragraph()
    separator = doc.add_paragraph("=" * 80)
    separator.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    # Add cover letter template section
    template_heading = doc.add_paragraph()
    template_heading.add_run("COVER LETTER TEMPLATE:").bold = True
    template_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    
    # Split template into paragraphs and add them
    template_paragraphs = template_text.split('\n')
    for para_text in template_paragraphs:
        if para_text.strip():  # Only add non-empty paragraphs
            template_para = doc.add_paragraph(para_text.strip())
            template_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    return doc

# Main interface
col1, col2 = st.columns([1, 1])

with col1:
    st.markdown('<div class="section-header">📝 Cover Letter Template</div>', unsafe_allow_html=True)
    
    if DOCX_SUPPORT:
        template_file = st.file_uploader(
            "Upload Word document",
            type=['docx'],
            help="Upload your cover letter template",
            key="template_file"
        )
        
        template_text = ""
        if template_file is not None:
            template_text = read_docx_file(template_file)
            if template_text:
                st.success("✅ Template loaded")
                with st.expander("Preview", expanded=False):
                    st.text_area("", template_text, height=100, disabled=True, key="template_preview")
        else:
            template_text = st.text_area(
                "Or paste template here:",
                placeholder="Paste your cover letter template...",
                height=150,
                key="template_input"
            )
    else:
        template_text = st.text_area(
            "Paste template here:",
            placeholder="Paste your cover letter template...",
            height=150,
            key="template_input"
        )

    st.markdown('<div class="section-header">🎯 Language Options</div>', unsafe_allow_html=True)
    
    language_option = st.radio(
        "Choose style:",
        ["Professional Language", "Match Template Language", "Custom"],
        key="language_radio"
    )
    
    custom_prompt = ""
    if language_option == "Custom":
        custom_prompt = st.text_area(
            "Custom instructions:",
            placeholder="e.g., 'Use creative but professional language'",
            height=60,
            key="custom_prompt_input"
        )

with col2:
    st.markdown('<div class="section-header">📋 Resume Information</div>', unsafe_allow_html=True)
    
    # Determine file types based on available support
    if DOCX_SUPPORT and PDF_SUPPORT:
        file_types = ['docx', 'pdf']
        help_text = "Upload your resume (.docx or .pdf format)"
        uploader_label = "Upload Word document or PDF"
    elif DOCX_SUPPORT:
        file_types = ['docx']
        help_text = "Upload your resume (.docx format)"
        uploader_label = "Upload Word document"
    else:
        file_types = []
        help_text = "File upload not available"
        uploader_label = "File upload disabled"
    
    resume_text = ""
    if file_types:
        resume_file = st.file_uploader(
            uploader_label,
            type=file_types,
            help=help_text,
            key="resume_uploader"
        )
        
        if resume_file is not None:
            if PDF_SUPPORT and resume_file.type == "application/pdf":
                resume_text = read_pdf_file(resume_file)
            else:
                resume_text = read_docx_file(resume_file)
            
            if resume_text:
                st.success("✅ Resume loaded")
                with st.expander("Preview", expanded=False):
                    st.text_area("", resume_text, height=100, disabled=True, key="resume_preview")
    
    if not resume_text:
        resume_text = st.text_area(
            "Paste resume info here:",
            placeholder="Paste key resume information...",
            height=150,
            key="resume_input"
        )

    st.markdown('<div class="section-header">💼 Job Description</div>', unsafe_allow_html=True)
    
    job_description = st.text_area(
        "Paste job description:",
        placeholder="Paste the complete job description here...",
        height=180,
        key="job_desc_input"
    )

# Options section
st.markdown('<div class="section-header">⚙️ Options</div>', unsafe_allow_html=True)

keep_one_page = st.checkbox(
    "Keep it one page",
    value=True,
    help="Instruct AI to keep cover letter to one page",
    key="keep_one_page"
)

# Generate default prompt when options change
st.session_state.default_prompt = generate_default_prompt(
    language_option, custom_prompt, keep_one_page
)

# Editable prompt section
st.markdown('<div class="section-header">✏️ Edit AI Prompt</div>', unsafe_allow_html=True)

edited_prompt = st.text_area(
    "Customize the AI instructions:",
    value=st.session_state.default_prompt,
    height=80,
    key="prompt_editor"
)

# Generate button and results
st.markdown('<div class="section-header">🚀 Generate</div>', unsafe_allow_html=True)

if st.button("📄 Generate Word Document", type="primary", use_container_width=True):
    if not DOCX_SUPPORT:
        st.error("❌ Cannot generate Word documents - python-docx dependency missing")
    elif not template_text:
        st.error("❌ Please provide a cover letter template")
    elif not resume_text:
        st.error("❌ Please provide resume information")
    elif not job_description:
        st.error("❌ Please provide the job description")
    else:
        with st.spinner("Creating document..."):
            st.session_state.generated_document = create_word_document(
                edited_prompt, job_description, template_text, resume_text
            )
        if st.session_state.generated_document:
            st.success("✅ Document generated!")

# Display download option
if st.session_state.generated_document:
    st.markdown("---")
    
    # Save document to bytes
    doc_buffer = io.BytesIO()
    st.session_state.generated_document.save(doc_buffer)
    doc_buffer.seek(0)
    
    # Generate filename
    doc_filename = f"cover_letter_AI_ready_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    st.download_button(
        label="💾 Download Word Document",
        data=doc_buffer.getvalue(),
        file_name=doc_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )
    
    st.info("📋 **Document contains:** AI prompt, resume info, job description, and your template. **Next:** Download → Copy all content → Paste into ChatGPT/Claude → AI will edit and remove instructions.")

# Footer
st.markdown("---")
st.write("**How to use:** Upload/paste template → Upload/paste resume → Paste job description → Generate → Download → Copy to AI tool")
st.write("**Pro tip:** The generated document has everything the AI needs - just copy/paste the entire content into ChatGPT or Claude.")
